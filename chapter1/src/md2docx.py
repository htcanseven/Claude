"""Convert the Chapter 1 markdown sources to Word documents with embedded figures."""
import re, sys, os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

FIGDIR = 'figures/'

# figure number -> image file, per document version
MAP_FULL = {
    '1.1': 'fig_stress_vs_tipspeed.png',
    '1.2': 'fig_classification_map.png',
    '1.3': 'fig_bending_modes.png',
    '1.4': 'fig_critical_speed_divergence.png',
    '1.5': 'fig_scaling_paths.png',
    '1.6': 'fig_converter_gate.png',
    '1.7': 'fig_geared_vs_directdrive.png',
    '1.8': 'fig_design_space.png',
}
MAP_CONS = {
    '1.1': 'fig_classification_map.png',
    '1.2': 'fig_geared_vs_directdrive.png',
    '1.3': 'fig_stress_vs_tipspeed.png',
    '1.4': 'fig_bending_modes.png',
    '1.5': 'fig_taylor_vortices.png',
    '1.6': 'fig_converter_gate.png',
    '1.7': 'fig_scaling_paths.png',
    '1.8': 'fig_critical_speed_divergence.png',
    '1.9': 'fig_design_space.png',
}
MAP_MIN = {
    '1.1': 'fig_stress_vs_tipspeed.png',
    '1.2': 'fig_classification_map.png',
    '1.3': 'fig_geared_vs_directdrive.png',
    '1.4': 'fig_bending_modes.png',
    '1.5': 'fig_converter_gate.png',
    '1.6': 'fig_critical_speed_divergence.png',
    '1.7': 'fig_scaling_paths.png',
    '1.8': 'fig_design_space.png',
}


def add_runs(par, text):
    """Render **bold** and *italic* inline markers into runs."""
    for tok in re.split(r'(\*\*.+?\*\*|\*[^*]+?\*)', text):
        if not tok:
            continue
        if tok.startswith('**') and tok.endswith('**'):
            par.add_run(tok[2:-2]).bold = True
        elif tok.startswith('*') and tok.endswith('*'):
            par.add_run(tok[1:-1]).italic = True
        else:
            par.add_run(tok)


def setup(doc):
    st = doc.styles['Normal']
    st.font.name, st.font.size = 'Cambria', Pt(11)
    st.paragraph_format.space_after = Pt(8)
    st.paragraph_format.line_spacing = 1.15
    for name, size, colour in [('Heading 1', 18, 0x1a1a1a), ('Heading 2', 14, 0x1f4e79),
                               ('Heading 3', 12, 0x1f4e79), ('Heading 4', 11, 0x333333)]:
        s = doc.styles[name]
        s.font.name, s.font.size, s.font.bold = 'Cambria', Pt(size), True
        s.font.color.rgb = RGBColor((colour >> 16) & 255, (colour >> 8) & 255, colour & 255)


def caption(doc, text):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(4), Pt(14)
    r = p.add_run(text)
    r.italic, r.font.size = True, Pt(9)
    return p


def equation(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.space_before = p.paragraph_format.space_after = Pt(8)
    r = p.add_run(text.strip())
    r.font.name, r.font.size = 'Cambria', Pt(11)
    r.italic = True
    return p


def add_table(doc, rows):
    header, body = rows[0], rows[1:]
    t = doc.add_table(rows=1, cols=len(header))
    t.style, t.alignment = 'Light Grid Accent 1', WD_TABLE_ALIGNMENT.CENTER
    for i, h in enumerate(header):
        cell = t.rows[0].cells[i]
        cell.text = ''
        add_runs(cell.paragraphs[0], h)
        for r in cell.paragraphs[0].runs:
            r.bold, r.font.size = True, Pt(9.5)
    for row in body:
        cells = t.add_row().cells
        for i, v in enumerate(row[:len(header)]):
            cells[i].text = ''
            add_runs(cells[i].paragraphs[0], v)
            for r in cells[i].paragraphs[0].runs:
                r.font.size = Pt(9.5)
    doc.add_paragraph().paragraph_format.space_after = Pt(6)
    return t


def convert(md_path, out_path, figmap, title):
    doc = Document()
    setup(doc)
    lines = open(md_path, encoding='utf-8').read().split('\n')
    i, pending_fig = 0, None

    while i < len(lines):
        ln = lines[i]
        stripped = ln.strip()

        # --- tables -------------------------------------------------------
        if stripped.startswith('|'):
            rows = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                cells = [c.strip() for c in lines[i].strip().strip('|').split('|')]
                if not all(re.fullmatch(r':?-{2,}:?', c) for c in cells):
                    rows.append(cells)
                i += 1
            if rows:
                add_table(doc, rows)
            continue

        # --- headings -----------------------------------------------------
        if stripped.startswith('#### '):
            doc.add_heading(stripped[5:], level=4)
        elif stripped.startswith('### '):
            doc.add_heading(stripped[4:], level=3)
        elif stripped.startswith('## '):
            doc.add_heading(stripped[3:], level=2)
        elif stripped.startswith('# '):
            doc.add_heading(stripped[2:], level=1)

        # --- figure anchor ------------------------------------------------
        elif re.match(r'\*\*\[Figure (1\.\d+)', stripped):
            num = re.match(r'\*\*\[Figure (1\.\d+)', stripped).group(1)
            img = figmap.get(num)
            if img and os.path.exists(FIGDIR + img):
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                p.add_run().add_picture(FIGDIR + img, width=Inches(5.4))
            else:
                p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run(f'[ Figure {num} — artwork to be prepared ]')
                r.italic, r.font.size = True, Pt(10)
                r.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
            pending_fig = num

        # --- figure caption ------------------------------------------------
        elif stripped.startswith('*Figure 1.') and stripped.endswith('*'):
            caption(doc, stripped.strip('*'))
            pending_fig = None

        # --- table caption --------------------------------------------------
        elif re.match(r'\*\*Table 1\.\d+\*\*', stripped):
            p = doc.add_paragraph()
            p.paragraph_format.space_before, p.paragraph_format.space_after = Pt(12), Pt(4)
            add_runs(p, stripped)
            for r in p.runs:
                r.font.size = Pt(9.5)

        # --- equations (indented four spaces) --------------------------------
        elif ln.startswith('    ') and stripped:
            equation(doc, ln)

        # --- horizontal rule --------------------------------------------------
        elif stripped == '---':
            doc.add_page_break()

        # --- bullets -----------------------------------------------------------
        elif stripped.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            add_runs(p, stripped[2:])

        # --- ordinary paragraph -------------------------------------------------
        elif stripped:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
            add_runs(p, stripped)

        i += 1

    doc.core_properties.title = title
    # python-docx's default template writes <w:zoom/> without the required
    # w:percent attribute, which fails schema validation; supply it.
    from docx.oxml.ns import qn
    zoom = doc.settings.element.find(qn('w:zoom'))
    if zoom is not None and zoom.get(qn('w:percent')) is None:
        zoom.set(qn('w:percent'), '100')
    doc.save(out_path)
    return out_path


if __name__ == '__main__':
    a = convert('ch1_full.md', 'Chapter_1_HighSpeed_FULL.docx', MAP_FULL,
                'High-Speed Electrical Machines — Chapter 1 (full version)')
    b = convert('ch1_minimal.md', 'Chapter_1_HighSpeed_MINIMAL.docx', MAP_MIN,
                'High-Speed Electrical Machines — Chapter 1 (minimal version)')
    c = convert('ch1_conservative.md', 'Chapter_1_HighSpeed_CONSERVATIVE.docx', MAP_CONS,
                'High-Speed Electrical Machines — Chapter 1 (conservative version)')
    for f in (a, b, c):
        print(f, f"{os.path.getsize(f)/1024:.0f} kB")
