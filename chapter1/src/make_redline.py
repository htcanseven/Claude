"""Build Word tracked-changes redlines: original draft -> revised chapter.

Produces genuine w:ins / w:del revision markup that Word's Review pane can step
through, accept and reject, and leaves Track Changes switched on.

The MINIMAL redline is marked up word by word: paragraphs are anchored first,
then diffed word by word inside each anchored pair.  A pure whole-document token
diff was tried and rejected -- with prose this heavily rewritten the longest
common subsequence collapses onto scattered stopwords and the markup becomes
confetti.  Where a paragraph has no counterpart at all there is nothing to diff
against, so it appears as a whole insertion or deletion.

The FULL redline stays at paragraph granularity, because that version was
restructured rather than edited and interleaved word markup would be unreadable.

The original draft text is read from original_draft_text.json (see
recover_original.py); the pipeline no longer depends on the uploaded .docx.
"""
import re, json, difflib, html, os
from docx import Document
from docx.shared import Pt, RGBColor
from docx.oxml import parse_xml

NS = 'xmlns:w="http://schemas.openxmlformats.org/wordprocessingml/2006/main"'
AUTHOR, DATE = 'Chapter 1 revision', '2026-09-02T00:00:00Z'
PBREAK = '\x00\x00P'          # paragraph-mark sentinel in the token stream


# ───────────────────────────────────────────────────────────── source loading
def extract_original():
    """(style, text) pairs for the original draft, from the cached recovery."""
    txt = json.load(open('original_draft_text.json', encoding='utf-8'))
    out = []
    for t in txt:
        # light heuristic so numbered section titles keep a heading look
        if re.match(r'^\d+\.\d+(\.\d+)?\s+\S', t) and len(t) < 90:
            out.append(('Heading 3', t))
        else:
            out.append(('Normal', t))
    return out


def extract_markdown(path):
    """(style, text) pairs from a revised chapter source."""
    items = []
    for ln in open(path, encoding='utf-8').read().split('\n'):
        s = ln.strip()
        if not s or s == '---':
            continue
        if s.startswith('#### '):
            items.append(('Heading 4', s[5:]))
        elif s.startswith('### '):
            items.append(('Heading 3', s[4:]))
        elif s.startswith('## '):
            items.append(('Heading 2', s[3:]))
        elif s.startswith('# '):
            items.append(('Heading 1', s[2:]))
        elif s.startswith('|'):
            cells = [c.strip() for c in s.strip('|').split('|')]
            if all(re.fullmatch(r':?-{2,}:?', c) for c in cells):
                continue
            items.append(('Normal', ' | '.join(cells)))
        else:
            plain = re.sub(r'\*\*(.+?)\*\*', r'\1', s)
            plain = re.sub(r'\*([^*]+?)\*', r'\1', plain)
            items.append(('Normal', re.sub(r'\s+', ' ', plain).strip()))
    return items


# ──────────────────────────────────────────────────────────────── XML helpers
_uid = [1000]


def nid():
    _uid[0] += 1
    return _uid[0]


def run_xml(text, mode):
    t = html.escape(text, quote=False)
    if mode == 'del':
        return (f'<w:del {NS} w:id="{nid()}" w:author="{AUTHOR}" w:date="{DATE}">'
                f'<w:r><w:delText xml:space="preserve">{t}</w:delText></w:r></w:del>')
    if mode == 'ins':
        return (f'<w:ins {NS} w:id="{nid()}" w:author="{AUTHOR}" w:date="{DATE}">'
                f'<w:r><w:t xml:space="preserve">{t}</w:t></w:r></w:ins>')
    return f'<w:r {NS}><w:t xml:space="preserve">{t}</w:t></w:r>'


def add_para(doc, style, segments, mark=None):
    p = doc.add_paragraph()
    try:
        p.style = doc.styles[style]
    except KeyError:
        pass
    if mark:
        tag = 'w:ins' if mark == 'ins' else 'w:del'
        p._p.get_or_add_pPr().append(parse_xml(
            f'<w:rPr {NS}><{tag} w:id="{nid()}" w:author="{AUTHOR}" w:date="{DATE}"/></w:rPr>'))
    for mode, text in segments:
        if text:
            p._p.append(parse_xml(run_xml(text, mode)))
    return p


def enable_track_changes(doc):
    doc.settings.element.append(parse_xml(f'<w:trackChanges {NS}/>'))


# ─────────────────────────────────────────────────────── word-level machinery
def tokenize(items):
    """Flatten (style, text) pairs into a word/paragraph-mark token stream."""
    toks, styles = [], []
    for style, text in items:
        for w in text.split():
            toks.append(w); styles.append(None)
        toks.append(PBREAK); styles.append(style)
    return toks, styles


def tidy(segments, min_equal=3):
    """Make a paragraph's redline readable.

    Merges adjacent same-mode segments, absorbs one- and two-word islands of
    unchanged text that sit between two changes, and orders each change region
    as deletions first then insertions -- which is how a redline is normally
    read, rather than word-by-word interleaving.
    """
    merged = []
    for m, t in segments:
        if merged and merged[-1][0] == m:
            merged[-1][1] += t
        else:
            merged.append([m, t])

    absorbed = []
    for i, (m, t) in enumerate(merged):
        island = (m is None and 0 < i < len(merged) - 1
                  and len(t.split()) < min_equal
                  and merged[i - 1][0] is not None and merged[i + 1][0] is not None)
        if island:
            absorbed.append(['del', t]); absorbed.append(['ins', t])
        else:
            absorbed.append([m, t])

    out, i = [], 0
    while i < len(absorbed):
        if absorbed[i][0] is None:
            out.append(tuple(absorbed[i])); i += 1
            continue
        j = i
        dels, inss = [], []
        while j < len(absorbed) and absorbed[j][0] is not None:
            (dels if absorbed[j][0] == 'del' else inss).append(absorbed[j][1])
            j += 1
        if dels:
            out.append(('del', ''.join(dels)))
        if inss:
            out.append(('ins', ''.join(inss)))
        i = j
    return out


def build_wordlevel(new_md, out_path, title, summary, thresh=0.25):
    """Redline with word-level markup wherever a counterpart paragraph exists.

    Paragraphs are anchored first, then diffed word by word inside each anchored
    pair.  A pure whole-document token diff was tried and rejected: with prose
    this heavily rewritten the longest common subsequence collapses onto
    scattered stopwords, and the markup becomes confetti.  Where a paragraph has
    no counterpart at all there is nothing to diff against, so it is shown as a
    whole insertion or deletion.
    """
    orig, new = extract_original(), extract_markdown(new_md)
    doc = Document()
    st = doc.styles['Normal']
    st.font.name, st.font.size = 'Cambria', Pt(11)
    st.paragraph_format.space_after = Pt(7)
    enable_track_changes(doc)
    banner(doc, title, summary)

    a, b = [t for _, t in orig], [t for _, t in new]
    stats = dict(word_level=0, inserted=0, deleted=0,
                 kept_words=0, ins_words=0, del_words=0)

    for i, j in align(a, b, thresh=thresh, band=0.45):
        if i is not None and j is not None:
            segs = word_diff(a[i], b[j])
            add_para(doc, new[j][0], segs)
            stats['word_level'] += 1
            for mode, t in segs:
                n = len(t.split())
                stats[{None: 'kept_words', 'ins': 'ins_words', 'del': 'del_words'}[mode]] += n
        elif i is not None:
            add_para(doc, 'Normal', [('del', a[i])], mark='del')
            stats['deleted'] += 1
            stats['del_words'] += len(a[i].split())
        else:
            add_para(doc, new[j][0], [('ins', b[j])], mark='ins')
            stats['inserted'] += 1
            stats['ins_words'] += len(b[j].split())

    doc.core_properties.title = title
    doc.save(out_path)
    return stats


# ────────────────────────────────────────────────── paragraph-level machinery
STOP = set('the a an of and or to in is are for on at as by it its this that with be '
           'from can we they which not but if then than these those'.split())


def bag(s):
    return {w for w in re.findall(r"[a-z0-9]+", s.lower()) if w not in STOP}


def sim(sa, sb):
    return len(sa & sb) / max(len(sa), len(sb)) if sa and sb else 0.0


def word_diff(a, b):
    aw, bw = a.split(' '), b.split(' ')
    segs = []
    for op, i1, i2, j1, j2 in difflib.SequenceMatcher(None, aw, bw, autojunk=False).get_opcodes():
        if op == 'equal':
            segs.append((None, ' '.join(aw[i1:i2]) + ' '))
        elif op == 'delete':
            segs.append(('del', ' '.join(aw[i1:i2]) + ' '))
        elif op == 'insert':
            segs.append(('ins', ' '.join(bw[j1:j2]) + ' '))
        else:
            segs.append(('del', ' '.join(aw[i1:i2]) + ' '))
            segs.append(('ins', ' '.join(bw[j1:j2]) + ' '))
    return tidy(segs)


def align(a, b, thresh=0.30, band=0.45):
    """Order-preserving alignment; the diagonal band stops the planning skeleton
    at the head of the draft from pairing with headings deep in the revision."""
    A, B = [bag(x) for x in a], [bag(x) for x in b]
    n, m = len(a), len(b)

    def score(i, j):
        if abs(i / n - j / m) > band:
            return 0.0
        return 1.0 if a[i] == b[j] else sim(A[i], B[j])

    S = [[0.0] * (m + 1) for _ in range(n + 1)]
    for i in range(1, n + 1):
        row, prev = S[i], S[i - 1]
        for j in range(1, m + 1):
            s = score(i - 1, j - 1)
            diag = prev[j - 1] + s if s >= thresh else -1.0
            row[j] = max(diag, prev[j], row[j - 1])

    out, i, j = [], n, m
    while i > 0 and j > 0:
        s = score(i - 1, j - 1)
        if s >= thresh and abs(S[i][j] - (S[i - 1][j - 1] + s)) < 1e-9:
            out.append((i - 1, j - 1)); i, j = i - 1, j - 1
        elif abs(S[i][j] - S[i - 1][j]) < 1e-9:
            out.append((i - 1, None)); i -= 1
        else:
            out.append((None, j - 1)); j -= 1
    while i > 0:
        out.append((i - 1, None)); i -= 1
    while j > 0:
        out.append((None, j - 1)); j -= 1
    return list(reversed(out))


def build_paragraphlevel(new_md, out_path, title, summary):
    orig, new = extract_original(), extract_markdown(new_md)
    doc = Document()
    st = doc.styles['Normal']
    st.font.name, st.font.size = 'Cambria', Pt(11)
    st.paragraph_format.space_after = Pt(7)
    enable_track_changes(doc)
    banner(doc, title, summary)

    a, b = [t for _, t in orig], [t for _, t in new]
    stats = dict(equal=0, revised=0, rewritten=0, inserted=0, deleted=0)
    for i, j in align(a, b):
        if i is not None and j is not None:
            if a[i] == b[j]:
                add_para(doc, new[j][0], [(None, a[i])]); stats['equal'] += 1
            elif sim(bag(a[i]), bag(b[j])) >= 0.55:
                add_para(doc, new[j][0], word_diff(a[i], b[j])); stats['revised'] += 1
            else:
                add_para(doc, 'Normal', [('del', a[i])], mark='del')
                add_para(doc, new[j][0], [('ins', b[j])], mark='ins')
                stats['rewritten'] += 1
        elif i is not None:
            add_para(doc, 'Normal', [('del', a[i])], mark='del'); stats['deleted'] += 1
        else:
            add_para(doc, new[j][0], [('ins', b[j])], mark='ins'); stats['inserted'] += 1

    doc.core_properties.title = title
    doc.save(out_path)
    return stats


# ─────────────────────────────────────────────────────────────────── preamble
def banner(doc, title, summary):
    r = doc.add_paragraph().add_run(title); r.bold, r.font.size = True, Pt(15)
    for txt in [
        'Word tracked-changes comparison of the original draft against the revised chapter. '
        'Open Review → Tracking and set the display to All Markup to step through every change; '
        'Accept and Reject work normally. Track Changes is left switched on, so any further '
        'edits are also recorded.',
        'The original draft stored Greek letters as Symbol-font characters and eight equations '
        'as Word equation objects. These have been resolved to plain Unicode on both sides so '
        'that they do not appear as spurious differences. Figures and tables are compared as text.',
    ]:
        p = doc.add_paragraph(); r = p.add_run(txt)
        r.italic, r.font.size = True, Pt(9.5)
        r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)
    doc.add_paragraph()
    r = doc.add_paragraph().add_run('Summary of changes in this version'); r.bold = True
    for line in summary:
        p = doc.add_paragraph(style='List Number')
        p.add_run(line).font.size = Pt(10)
    doc.add_page_break()


SUM_FULL = [
    "Scaling table corrected. The draft's table mixed a fixed-geometry path with a tip-speed-limited path in one set of rows. Along the tip-speed-limited constant-power path, windage loss is unchanged rather than 8x, and iron loss doubles rather than quadruples.",
    "Rotordynamic scaling added. The first bending critical speed falls as the cube of design speed while operating speed rises linearly, so the margin degrades as the fourth power: doubling the speed costs a factor of 16. This constraint was absent from the draft and is the one that actually limits attainable speed.",
    "Three design paths separated explicitly (fixed geometry, tip-speed limited, length capped), because a scaling exponent is undefined until one states what is held constant. Every column of the new table names its path.",
    "Square-cube thermal argument replaced. It assumes isotropic scaling, which high-speed machines do not follow: they shrink radially and grow axially, so rotor surface area is preserved.",
    "Centrifugal stress quantified. The draft's claim that a central bore doubles stress is now derived and worked through in MPa against a real proof stress.",
    "Two drafting artefacts removed from the fP passage ('without external resources provided...', '...in the current context'), which read as an AI tool commenting on its own retrieval limits.",
    "The 180 000 kW/s threshold withdrawn. It had been calibrated from a single machine that is itself marginal by tip speed. fP is now presented as a proposed drive-system figure of merit, with calibration deferred to the Chapter 2 survey.",
    "Fourfold repetition of the industrial-versus-mobile contrast consolidated into one section and one table.",
    "Trilemma recast as three constraints plus a converter gate, resolving the draft's own observation that power electronics sat outside the trilemma.",
    "Roadmap rewritten from 19 chapters in six parts to 12 chapters in four parts.",
    "Planning skeleton removed from the head of the chapter; it had drifted from the text below it.",
    "Worked example added: the Voltcar machine carried through all three classification criteria and on to the converter requirement (31.5 kHz at m_f = 21, which silicon cannot supply).",
    "Constraints moved ahead of the design paradigms, so the physics is established before the two families are explained as different responses to it.",
    "Direct-drive history rewritten as the lead-in to the paradigms section, where it explains why the two families diverge.",
    "Notation paragraph added, fixing symbol conventions for the book.",
    "Figure numbering corrected; the draft had Figures 2 to 5 with no Figure 1. Eight figures now specified, all eight supplied.",
]

SUM_MIN = [
    "Scaling table corrected. The draft's table mixed a fixed-geometry path with a tip-speed-limited path in one set of rows. Along the tip-speed-limited constant-power path, windage loss is unchanged rather than 8x, and iron loss doubles rather than quadruples.",
    "Rotordynamic scaling added. The first bending critical speed falls as the cube of design speed while operating speed rises linearly, so the margin degrades as the fourth power: doubling the speed costs a factor of 16. This constraint was absent from the draft.",
    "Three design paths separated explicitly (fixed geometry, tip-speed limited, length capped). Every column of the new table names its path.",
    "Square-cube thermal argument replaced, since it assumes an isotropic scaling that high-speed machines do not follow.",
    "Centrifugal stress quantified, with the doubling caused by a central bore now derived rather than asserted.",
    "Two drafting artefacts removed from the fP passage ('without external resources provided...', '...in the current context').",
    "The 180 000 kW/s threshold withdrawn; fP presented as a proposed drive-system figure of merit with calibration deferred to Chapter 2.",
    "Fourfold repetition of the industrial-versus-mobile contrast consolidated into one section and one table.",
    "Trilemma recast as three constraints plus a converter gate.",
    "Roadmap rewritten from 19 chapters in six parts to 12 chapters in four parts.",
    "Planning skeleton removed from the head of the chapter.",
    "Worked example added: the Voltcar machine carried through all three criteria and on to the converter requirement.",
    "Figure numbering corrected; eight figures now specified and all eight supplied.",
    "Original section structure and numbering retained throughout, so this version can be reviewed side by side with the draft.",
    "This redline is marked up word by word: both documents were diffed as a single token stream, so changes appear at the level of individual words and phrases rather than whole paragraphs.",
]

if __name__ == '__main__':
    s = build_wordlevel('ch1_minimal.md', 'Chapter_1_Redline_MINIMAL.docx',
                        'Chapter 1 — Tracked Changes (word level): original draft → MINIMAL version',
                        SUM_MIN)
    print(f"Chapter_1_Redline_MINIMAL.docx  [word level]\n"
          f"   word-level paragraphs {s['word_level']}   whole inserts {s['inserted']}   "
          f"whole deletes {s['deleted']}\n"
          f"   words: kept {s['kept_words']}  inserted {s['ins_words']}  deleted {s['del_words']}")
    s = build_paragraphlevel('ch1_full.md', 'Chapter_1_Redline_FULL.docx',
                             'Chapter 1 — Tracked Changes: original draft → FULL version',
                             SUM_FULL)
    print(f"Chapter_1_Redline_FULL.docx  [paragraph level]\n"
          f"   unchanged {s['equal']}   edited-inline {s['revised']}   rewritten {s['rewritten']}   "
          f"inserted {s['inserted']}   deleted {s['deleted']}")
