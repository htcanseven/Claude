"""Rebuild the original draft as a Word document from the recovered text.

The uploaded .docx was lost when the container was recycled.  Its text was
recovered losslessly from the tracked-changes redlines (reject-all reproduces
the source document exactly) and cached in original_draft_text.json.

This reconstruction carries the text only.  The four images embedded in the
draft, its Word comments, and its original character formatting are not
recoverable and are not present here.
"""
import json, re
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

# paragraph indices that carried a heading style in the draft
HEADINGS = {
    0, 1, 5, 7, 8, 13, 18, 19, 22, 25, 27, 28, 30, 32, 33, 35, 37,
    40, 48, 50, 54, 66, 68, 72, 77, 80, 82, 85, 88, 93, 95, 100, 102, 104, 107,
    112, 116, 118, 120, 125, 131, 134, 139, 141, 142, 144, 147, 153, 156, 162,
    170, 172, 177, 185, 187, 194, 197, 210, 212, 218, 223, 229, 242, 246, 252,
    254, 256, 258, 260, 262, 264,
}
EQUATIONS = {43, 179, 220}
TABLE_CAPTIONS = {60, 203, 234}
TABLE_ROWS = set(range(61, 65)) | set(range(204, 210)) | set(range(235, 241))
FIG_PLACEHOLDER = {70, 71, 155, 175, 176, 244, 245}

paras = json.load(open('original_draft_text.json', encoding='utf-8'))
doc = Document()
st = doc.styles['Normal']
st.font.name, st.font.size = 'Cambria', Pt(11)
st.paragraph_format.space_after = Pt(8)
for nm, sz in [('Heading 1', 16), ('Heading 2', 14), ('Heading 3', 12)]:
    s = doc.styles[nm]
    s.font.name, s.font.size, s.font.bold = 'Cambria', Pt(sz), True
    s.font.color.rgb = RGBColor(0x1a, 0x1a, 0x1a)

r = doc.add_paragraph().add_run('The first Ch of HS machines — original draft (text reconstruction)')
r.bold, r.font.size = True, Pt(14)
note = doc.add_paragraph().add_run(
    'Text of the original draft, reproduced exactly. Recovered from the tracked-changes '
    'redlines, whose reject-all output reproduces the source document verbatim; verified '
    'identical from two independent redlines. Text only: the four embedded images, the Word '
    'comments and the original character formatting are not part of this reconstruction, and '
    'the figure placeholders appear below exactly as they read in the draft.')
note.italic, note.font.size = True, Pt(9.5)
note.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
doc.add_paragraph()

i = 0
while i < len(paras):
    p = paras[i]
    if i in TABLE_ROWS:
        rows = []
        while i < len(paras) and i in TABLE_ROWS:
            rows.append([c.strip() for c in paras[i].split('|')]); i += 1
        w = max(len(x) for x in rows)
        t = doc.add_table(rows=0, cols=w); t.style = 'Table Grid'
        for row in rows:
            cells = t.add_row().cells
            for k, v in enumerate(row + [''] * (w - len(row))):
                cells[k].text = ''
                run = cells[k].paragraphs[0].add_run(v); run.font.size = Pt(9.5)
        doc.add_paragraph()
        continue

    if i in FIG_PLACEHOLDER:
        q = doc.add_paragraph(); q.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = q.add_run(p if p != '[FIGURE]' else '[ image in the original draft ]')
        run.italic, run.font.size = True, Pt(10)
        run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    elif i in HEADINGS:
        lvl = 3 if re.match(r'^\d+\.\d+\.\d+', p) else (2 if re.match(r'^\d+\.\d+\.?\s', p) else 3)
        doc.add_heading(p, level=lvl)
    elif i in EQUATIONS:
        q = doc.add_paragraph(); q.paragraph_format.left_indent = Pt(36)
        run = q.add_run(p); run.italic = True
    elif i in TABLE_CAPTIONS:
        q = doc.add_paragraph(); run = q.add_run(p); run.bold, run.font.size = True, Pt(9.5)
    else:
        doc.add_paragraph(p).alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    i += 1

doc.core_properties.title = 'The first Ch of HS machines — original draft (text reconstruction)'
doc.save('Chapter_1_ORIGINAL_DRAFT_reconstruction.docx')
print('wrote Chapter_1_ORIGINAL_DRAFT_reconstruction.docx')
print(f'  {len(paras)} paragraphs, {sum(len(x.split()) for x in paras):,} words')
